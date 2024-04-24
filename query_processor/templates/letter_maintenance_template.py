import uuid
from datetime import datetime

from docx import Document
from docx2pdf import convert as convert_docx_to_pdf

from query_processor.templates.template import Template


class LetterMaintenanceDocxTemplate(Template):
    __fields = [
        'date', 'id',  'timestamp', 'main_application',
        'patent_id', 'patent_name', 'payment_order', 'payment_date', 'payment_count'
    ]

    template = './templates/docs/LetterMaintenance.docx'

    def get_fields(self):
        return self.__fields

    def fill(self, fields) -> str:
        doc = Document(self.template)
        output_file = f"./out/Letter_maintenance_{uuid.uuid4()}.docx"
        for paragraph in doc.paragraphs:
            for key, value in fields.items():
                for run in paragraph.runs:
                    if f'{{{{ {key} }}}}' in run.text:
                        if isinstance(value, datetime):
                            value = value.strftime('%d.%m.%Y')
                        run.text = run.text.replace(f'{{{{ {key} }}}}', str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in fields.items():
                            for run in paragraph.runs:
                                if f'{{{{ {key} }}}}' in run.text:
                                    if isinstance(value, datetime):
                                        value = value.strftime('%d.%m.%Y')
                                    run.text = run.text.replace(f'{{{{ {key} }}}}', str(value))

        doc.save(output_file)
        return output_file


class LetterMaintenancePdfTemplate(LetterMaintenanceDocxTemplate):

    # Для конвертации word в pdf используется библиотека docx2pdf
    # также есть вариант использовать pandoc
    # TODO: попробовать другие варианты конвертации
    def fill(self, fields) -> str:
        docx_file = super().fill(fields)
        pdf_file = docx_file.replace('.docx', '.pdf')
        convert_docx_to_pdf(docx_file, pdf_file)

        return pdf_file
